import os
import json
import re
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import argparse
import glob
from typing import Dict, List, Tuple

def clean_tag(text: str) -> str:
    """Clean and standardize tag format."""
    text = text.strip()
    # Remove numbering prefix
    text = re.sub(r'^\d+\.\s*', '', text)
    # Remove trailing punctuation
    text = re.sub(r'[.,;:!?]+$', '', text)
    # Remove extra whitespace
    text = ' '.join(text.split())
    return text

def is_tag(text: str, para) -> Tuple[bool, float]:
    """Enhanced tag validation with confidence score."""
    text = clean_tag(text)
    confidence = 0.0
    
    # Length check (10-50 chars)
    if 10 <= len(text) <= 50:
        confidence += 0.2
    else:
        return False, 0.0
        
    # Format checks
    if text[0].isupper():  # Starts with capital
        confidence += 0.2
    if not text.endswith(('.', '!', '?')):  # No sentence-ending punctuation
        confidence += 0.2
    if not re.search(r'[A-Z]{3,}', text):  # No long uppercase sequences
        confidence += 0.2
        
    # Formatting check
    if para.runs:
        has_formatting = any(run.bold for run in para.runs) or any(run.underline for run in para.runs)
        if has_formatting:
            confidence += 0.2
            
    # Pattern checks
    tag_patterns = [
        r'^[A-Z][a-z]+(\s+[A-Z][a-z]+)*$',  # Title Case
        r'^[A-Z][a-z]+(\s+[A-Z][a-z]+)*:',  # Title Case with colon
        r'^[A-Z][a-z]+(\s+[A-Z][a-z]+)*\?$',  # Title Case with question mark
    ]
    if any(re.match(pattern, text) for pattern in tag_patterns):
        confidence += 0.2
        
    return confidence >= 0.6, confidence

def extract_citation_features(text: str) -> Dict:
    """Extract features from citation text."""
    features = {
        'has_year': False,
        'year': None,
        'has_author': False,
        'author': None,
        'has_url': False,
        'has_publication': False,
        'confidence': 0.0
    }
    
    # Year check
    year_match = re.search(r'(19|20)\d{2}', text)
    if year_match:
        features['has_year'] = True
        features['year'] = year_match.group(0)
        features['confidence'] += 0.3
        
    # Author check (capitalized word before year)
    if year_match:
        before_year = text[:year_match.start()].strip()
        author_match = re.search(r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)', before_year)
        if author_match:
            features['has_author'] = True
            features['author'] = author_match.group(0)
            features['confidence'] += 0.3
            
    # URL/publication check
    if re.search(r'https?://|www\.', text):
        features['has_url'] = True
        features['confidence'] += 0.2
    if re.search(r'\b[A-Z][a-z]+\b.*(19|20)\d{2}', text):
        features['has_publication'] = True
        features['confidence'] += 0.2
        
    return features

def is_citation(text: str) -> Tuple[bool, float]:
    """Enhanced citation validation with confidence score."""
    text = text.strip()
    if len(text) > 300 or text.isupper():
        return False, 0.0
        
    features = extract_citation_features(text)
    return features['confidence'] >= 0.6, features

def format_run(run) -> str:
    """Format text run with HTML tags."""
    text = run.text
    if not text:
        return ""
        
    # Track formatting features
    formatting = []
    if run.bold:
        formatting.append('b')
    if run.underline:
        formatting.append('u')
    if run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
        formatting.append('hl')
    if run.italic:
        formatting.append('i')
        
    # Apply formatting
    for fmt in formatting:
        text = f"<{fmt}>{text}</{fmt}>"
    return text

def collect_fulltext(para, current_fulltext: List[str], current_markup: List[str]) -> Tuple[List[str], List[str]]:
    """Collect fulltext with formatting."""
    text = para.text.strip()
    if not text:
        return current_fulltext, current_markup
        
    # Check for evidence end markers
    end_markers = [
        r'^\d+\.\s*[A-Z]',
        r'^[A-Z][A-Z\s]+$',
        r'^[A-Z][a-z]+\s+\d{2,4}',
    ]
    
    if any(re.match(marker, text) for marker in end_markers):
        return current_fulltext, current_markup
        
    # Add to fulltext if it's evidence
    formatted_text = "".join([format_run(run) for run in para.runs])
    current_markup.append(formatted_text)
    current_fulltext.append(text)
    return current_fulltext, current_markup

def calculate_candidate_confidence(tag: str, citation: str, fulltext: str, tag_conf: float, citation_conf: float) -> float:
    """Calculate overall confidence score for candidate."""
    confidence = 0.0
    
    # Tag confidence
    confidence += tag_conf * 0.3
    
    # Citation confidence
    confidence += citation_conf * 0.3
    
    # Fulltext quality
    if len(fulltext) >= 200:
        confidence += 0.2
    if '.' in fulltext:
        confidence += 0.1
        
    # Formatting check
    if '<b>' in fulltext or '<u>' in fulltext:
        confidence += 0.1
        
    return confidence

def is_heading_or_boilerplate(text: str) -> bool:
    """Check if text is a heading or boilerplate."""
    text = text.strip()
    if text.isupper() and len(text) < 20:
        return True
    if re.match(r'^page \d+$', text, re.IGNORECASE):
        return True
    if re.match(r'^contention \d+', text, re.IGNORECASE):
        return True
    return False

def extract_candidates_from_docx(docx_dir: str, output_file: str):
    """Extract evidence candidates from DOCX files with enhanced validation."""
    candidates = []
    docx_files = glob.glob(os.path.join(docx_dir, '**', '*.docx'), recursive=True)
    
    for filepath in docx_files:
        filename = os.path.basename(filepath)
        if filename.startswith('~$'):
            continue
        print(f"Processing {filepath}...")
        
        doc = Document(filepath)
        current_tag = None
        current_citation = None
        current_fulltext = []
        current_markup = []
        tag_confidence = 0.0
        citation_confidence = 0.0
        
        i = 0
        while i < len(doc.paragraphs):
            para = doc.paragraphs[i]
            text = para.text.strip()
            
            if not text:
                i += 1
                continue
                
            # Detect tag
            is_valid_tag, tag_conf = is_tag(text, para)
            if is_valid_tag:
                current_tag = clean_tag(text)
                tag_confidence = tag_conf
                current_citation = None
                current_fulltext = []
                current_markup = []
                i += 1
                continue
                
            # Detect citation
            is_valid_citation, citation_features = is_citation(text)
            if is_valid_citation:
                if not current_tag:
                    print(f"[SKIP] Citation found without preceding tag at paragraph {i}: {text[:60]}...")
                    i += 1
                    continue
                    
                current_citation = text
                citation_confidence = citation_features['confidence']
                current_fulltext = []
                current_markup = []
                i += 1
                
                # Collect fulltext
                while i < len(doc.paragraphs):
                    para_ft = doc.paragraphs[i]
                    text_ft = para_ft.text.strip()
                    
                    if not text_ft or len(text_ft) < 20 or is_heading_or_boilerplate(text_ft):
                        i += 1
                        continue
                        
                    if is_tag(text_ft, para_ft)[0] or is_citation(text_ft)[0]:
                        break
                        
                    current_fulltext, current_markup = collect_fulltext(para_ft, current_fulltext, current_markup)
                    i += 1
                    
                # Create candidate
                fulltext_str = "\n".join(current_fulltext).strip()
                confidence = calculate_candidate_confidence(
                    current_tag, current_citation, fulltext_str,
                    tag_confidence, citation_confidence
                )
                
                if confidence >= 0.6:  # Minimum confidence threshold
                    candidate = {
                        "tag": current_tag,
                        "citation": current_citation,
                        "fulltext": fulltext_str,
                        "preserved_markup": "\n".join(current_markup),
                        "confidence": confidence,
                        "features": {
                            "tag_confidence": tag_confidence,
                            "citation_confidence": citation_confidence,
                            "citation_features": citation_features,
                            "fulltext_length": len(fulltext_str),
                            "has_formatting": '<b>' in fulltext_str or '<u>' in fulltext_str
                        }
                    }
                    candidates.append(candidate)
                else:
                    print(f"[SKIP] Low-confidence candidate: tag='{current_tag[:30]}', citation='{current_citation[:30]}', confidence={confidence:.2f}")
                    
                current_tag = None
                current_citation = None
                current_fulltext = []
                current_markup = []
                continue
                
            i += 1
            
    with open(output_file, 'w') as f:
        for candidate in candidates:
            f.write(json.dumps(candidate) + '\n')
            
    print(f"Extracted {len(candidates)} candidates to {output_file}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument('--docx_dir', type=str, default='converted_docx', help='Directory with DOCX files')
    parser.add_argument('--output_file', type=str, default='candidates.jsonl', help='Output JSONL file')
    args = parser.parse_args()
    extract_candidates_from_docx(args.docx_dir, args.output_file) 